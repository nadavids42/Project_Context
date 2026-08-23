"""DOCX parser: paragraphs, lists, and table cells in deterministic
reading order (Section 8/15).

python-docx's own `.paragraphs` and `.tables` collections are separate
top-level lists that lose interleaving order — they can't tell you a
table appeared *between* two particular paragraphs. `_iter_block_items`
is the standard recipe for recovering true reading order by walking the
document body's direct XML children instead.

Documented limitation: headers and footers are not extracted. They live
outside the body, are scoped per section rather than per paragraph, and
are frequently boilerplate (page numbers, running titles) rather than
project-material text.
"""

from __future__ import annotations

import io
from collections.abc import Iterator

import docx
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from project_context.domain.evidence import ParseStatus
from project_context.parsers.models import ParseResult, TextBlock

PARSER_NAME = "docx"
PARSER_VERSION = "1"


def _iter_block_items(document: DocxDocument) -> Iterator[Paragraph | Table]:
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def _is_list_paragraph(paragraph: Paragraph) -> bool:
    style_name = paragraph.style.name if paragraph.style else ""
    return bool(style_name) and "list" in style_name.lower()


def parse_docx(data: bytes) -> ParseResult:
    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:  # python-docx raises assorted exceptions for a bad zip/XML
        return ParseResult(
            status=ParseStatus.FAILED,
            parser_name=PARSER_NAME,
            parser_version=PARSER_VERSION,
            normalized_text="",
            warnings=(f"could not open DOCX: {exc}",),
        )

    blocks: list[TextBlock] = []
    text_parts: list[str] = []
    cursor = 0
    paragraph_ordinal = 0
    table_ordinal = 0

    def append_block(text: str, section_path: str) -> None:
        nonlocal cursor
        if not text.strip():
            return
        start = cursor
        text_parts.append(text)
        cursor += len(text)
        blocks.append(
            TextBlock(text=text, char_start=start, char_end=cursor, section_path=section_path)
        )
        text_parts.append("\n\n")
        cursor += 2

    for item in _iter_block_items(document):
        if isinstance(item, Paragraph):
            paragraph_ordinal += 1
            kind = "list item" if _is_list_paragraph(item) else "paragraph"
            append_block(item.text, f"{kind} {paragraph_ordinal}")
        elif isinstance(item, Table):
            table_ordinal += 1
            for row_idx, row in enumerate(item.rows, start=1):
                for col_idx, cell in enumerate(row.cells, start=1):
                    append_block(cell.text, f"table {table_ordinal} row {row_idx} col {col_idx}")

    if not blocks:
        return ParseResult(
            status=ParseStatus.EMPTY,
            parser_name=PARSER_NAME,
            parser_version=PARSER_VERSION,
            normalized_text="",
        )

    normalized_text = "".join(text_parts).strip()
    return ParseResult(
        status=ParseStatus.PARSED,
        parser_name=PARSER_NAME,
        parser_version=PARSER_VERSION,
        normalized_text=normalized_text,
        blocks=tuple(blocks),
    )
