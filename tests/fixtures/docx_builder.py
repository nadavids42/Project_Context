"""Small python-docx-based DOCX fixture builder — real library output,
not a hand-rolled binary, so fixtures stay trivially editable."""

from __future__ import annotations

import io

from docx import Document


def build_docx_with_paragraphs_and_table() -> bytes:
    document = Document()
    document.add_paragraph("First paragraph.")
    document.add_paragraph("Bullet one", style="List Bullet")
    document.add_paragraph("Bullet two", style="List Bullet")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A1"
    table.cell(0, 1).text = "B1"
    table.cell(1, 0).text = "A2"
    table.cell(1, 1).text = "B2"
    document.add_paragraph("Last paragraph.")

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def build_empty_docx() -> bytes:
    buf = io.BytesIO()
    Document().save(buf)
    return buf.getvalue()
