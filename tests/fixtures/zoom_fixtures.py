"""Representative Zoom-to-Drive export fixtures (Section 11.6; Prompt
13). These are hand-built to match Zoom's own documented default
export shapes for cloud-recording transcripts and an AI Companion
meeting summary — not pulled from a live Zoom account (Section 15:
"tests must not require live credentials"). See
`tests/unit/test_zoom_drive_compatibility.py` for the compatibility
finding this fixture surfaces: Zoom's own VTT transcript export does
**not** use WebVTT `<v Speaker>` voice tags — it puts the speaker name
as a plain-text `"Speaker Name: "` prefix inside each cue's own text.
`project_context.parsers.vtt_parser` only recognizes `<v>` tags for
speaker extraction/merging, so this fixture, run through it, correctly
parses and stores every word of the transcript (nothing is lost, and it
remains fully evidence-linkable/citable) but loses per-speaker turn
boundaries — every cue merges into one block since none of them carry a
distinct extracted `speaker` value. That is a real, documented
limitation this prompt discovered and reports rather than silently
"fixes" inside a connector/compatibility prompt — see the Prompt 13
report for the recommendation.
"""

from __future__ import annotations

import io

from docx import Document

#: Filename following Zoom's own default cloud-recording naming
#: convention: `GMT` + `YYYYMMDD-HHMMSS` + the recording's topic +
#: `.transcript.vtt`.
ZOOM_VTT_FILENAME = "GMT20260601-140000_Acme-Kickoff.transcript.vtt"

#: Zoom's real `audio_transcript.vtt` shape: numbered cues (a bare
#: integer line before each timestamp line — WebVTT's optional cue
#: identifier), no `<v>` voice tags, speaker name as a plain-text
#: `"Name: "` prefix inside the cue text itself.
ZOOM_VTT_TRANSCRIPT_BYTES = b"""WEBVTT

1
00:00:00.000 --> 00:00:03.500
Alice Smith: Thanks everyone for joining today's kickoff.

2
00:00:03.500 --> 00:00:07.800
Bob Jones: Happy to be here, excited to get started.

3
00:00:07.800 --> 00:00:12.000
Alice Smith: We decided to ship the Acme rollout in July.
"""

#: A Zoom in-meeting chat export, Zoom's own default filename suffix.
ZOOM_CHAT_FILENAME = "GMT20260601-140000_Acme-Kickoff.chat.txt"
ZOOM_CHAT_BYTES = (
    b"00:02:15\tBob Jones:\tI'll drop the doc link here.\n"
    b"00:02:20\tBob Jones:\thttps://example.com/doc\n"
)

#: Zoom AI Companion's manually-saved meeting-summary export — no fixed
#: filename standard from Zoom itself, so this uses a plausible
#: human-chosen name rather than a documented default.
ZOOM_SUMMARY_TXT_FILENAME = "Meeting Summary - Acme Kickoff - 2026-06-01.txt"
ZOOM_SUMMARY_TXT_TEXT = """Quick recap
Alice and Bob kicked off the Acme rollout project and agreed on July as the target ship date.

Summary
The team reviewed the current project timeline. Alice confirmed the Acme rollout will ship in \
July. Bob will send the requirements doc by Friday.

Next steps
Bob to send the requirements doc by Friday.
Alice to schedule the next check-in.
"""

ZOOM_SUMMARY_DOCX_FILENAME = "Meeting Summary - Acme Kickoff - 2026-06-01.docx"


def build_zoom_summary_docx() -> bytes:
    """The same summary content as `ZOOM_SUMMARY_TXT_TEXT`, as a real
    `.docx` (python-docx output, not a hand-rolled binary) — Zoom AI
    Companion summaries are commonly copied into a Word document before
    landing in Drive."""
    document = Document()
    document.add_heading("Quick recap", level=2)
    document.add_paragraph(
        "Alice and Bob kicked off the Acme rollout project and agreed on July as the "
        "target ship date."
    )
    document.add_heading("Summary", level=2)
    document.add_paragraph(
        "The team reviewed the current project timeline. Alice confirmed the Acme "
        "rollout will ship in July. Bob will send the requirements doc by Friday."
    )
    document.add_heading("Next steps", level=2)
    document.add_paragraph("Bob to send the requirements doc by Friday.", style="List Bullet")
    document.add_paragraph("Alice to schedule the next check-in.", style="List Bullet")

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()
